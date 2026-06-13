"""
Script to create test fixture files (DOCX templates and input case files)
for end-to-end testing of the Tax Investigation System.

Usage: python _create_fixtures.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from openpyxl import Workbook

BASE_DIR = Path.cwd()
SAMPLE_DIR = BASE_DIR / "sample"
INPUT_DIR = BASE_DIR / "input"
SAMPLE_DIR.mkdir(exist_ok=True)
INPUT_DIR.mkdir(exist_ok=True)


def create_docx(path: Path, paragraphs: list[str], bold_first: bool = False, table_data: list[list[str]] | None = None) -> None:
    """Create a simple .docx file with paragraphs and optional table."""
    doc = Document()
    for i, text in enumerate(paragraphs):
        p = doc.add_paragraph(text)
        if bold_first and i == 0:
            for run in p.runs:
                run.bold = True
            else:
                run = p.add_run(text)
                run.bold = True

    if table_data:
        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
        for i, row_data in enumerate(table_data):
            for j, cell_text in enumerate(row_data):
                table.rows[i].cells[j].text = cell_text

    doc.save(str(path))


def create_xlsx(path: Path, sheet_name: str, headers: list[str], rows: list[list[str]]) -> None:
    """Create a simple .xlsx file with a single sheet."""
    wb = Workbook()
    ws = wb.active
    ws.title = sheet_name
    ws.append(headers)
    for row in rows:
        ws.append(row)
    wb.save(str(path))


# ====== 1. Notice Template DOCX ======
create_docx(
    SAMPLE_DIR / "Notice_Template.docx",
    [
        "NOTICE UNDER SECTION 133(6) OF THE INCOME TAX ACT, 1961",
        "",
        "To,",
        "{{ assessee_name }}",
        "PAN: {{ pan }}",
        "Assessment Year: {{ assessment_year }}",
        "",
        "Case ID: {{ case_id }}",
        "Decision Type: {{ decision_type }}",
        "Reason Codes: {{ reason_codes }}",
        "",
        "Discrepancies Found:",
        "{{ discrepancies_block }}",
        "",
        "Generated at: {{ generated_at }}",
        "",
        "This is a computer-generated notice.",
    ],
)

# ====== 2. Report Template DOCX ======
create_docx(
    SAMPLE_DIR / "Tax_Investigation_Report_Template.docx",
    [
        "TAX INVESTIGATION REPORT",
        "",
        "Assessee Name: {{ assessee_name }}",
        "PAN: {{ pan }}",
        "Assessment Year: {{ assessment_year }}",
        "",
        "Case ID: {{ case_id }}",
        "Decision: {{ decision_type }}",
        "Reason Codes: {{ reason_codes }}",
        "",
        "Discrepancy Details:",
        "{{ discrepancies_block }}",
        "",
        "Report generated at: {{ generated_at }}",
    ],
)

# ====== 3. Form 16 test data (DOCX) ======
create_docx(
    INPUT_DIR / "Form16.docx",
    [
        "FORM 16 - SALARY CERTIFICATE",
        "Employer: ABC Corp Ltd",
        "TAN: TAN12345E",
        "Employee Name: Rajesh Kumar",
        "PAN: ABCPE1234F",
        "Assessment Year: 2025-26",
        "",
        "Gross Salary: 1,200,000",
        "TDS Deducted: 180,000",
        "Chapter VI-A Deductions: 50,000",
    ],
)

# ====== 4. AIS test data (XLSX - use openpyxl for structure) ======
create_xlsx(
    INPUT_DIR / "AIS.xlsx",
    "AIS Data",
    ["PAN", "Name", "AY", "FY"],
    [
        ["ABCPE1234F", "Rajesh Kumar", "2025-26", "2024-25"],
    ],
)

# Add more AIS data rows with amounts
wb = Workbook()
ws = wb.active
ws.title = "AIS Data"
ws.append(["PAN", "Name", "AY", "FY"])
ws.append(["ABCPE1234F", "Rajesh Kumar", "2025-26", "2024-25"])
ws.append([])
ws.append(["Category", "Amount"])
ws.append(["Salary", "1250000"])
ws.append(["Interest", "25000"])
ws.append(["Dividend", "5000"])
ws.append(["Securities", "10000"])
ws.append(["TDS", "185000"])
ws.append(["Bank Deposit", "800000"])
wb.save(str(INPUT_DIR / "AIS.xlsx"))

# ====== 5. ITR Extract test data (DOCX) ======
create_docx(
    INPUT_DIR / "ITR_Extract.docx",
    [
        "ITR EXTRACT - INCOME TAX RETURN",
        "Assessee Name: Rajesh Kumar",
        "PAN: ABCPE1234F",
        "Assessment Year: 2025-26",
        "",
        "Salary: 1200000",
        "Other Sources: 15000",
        "Interest: 20000",
        "Dividend: 5000",
        "Securities: 10000",
        "TDS: 180000",
        "Deductions: 50000",
        "Total Income: 1240000",
    ],
)

print("[OK] Test fixtures created successfully!")
print(f"  - Sample notice: {SAMPLE_DIR / 'Notice_Template.docx'}")
print(f"  - Sample report: {SAMPLE_DIR / 'Tax_Investigation_Report_Template.docx'}")
print(f"  - Input Form16:  {INPUT_DIR / 'Form16.docx'}")
print(f"  - Input AIS:     {INPUT_DIR / 'AIS.xlsx'}")
print(f"  - Input ITR:     {INPUT_DIR / 'ITR_Extract.docx'}")
